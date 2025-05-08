import os
import faiss
import pandas as pd
from docx import Document
from sentence_transformers import SentenceTransformer
import numpy as np
import openai
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
import tempfile
import shutil

# === Setup Azure OpenAI API ===
openai.api_type = "azure"
openai.api_key = "d004ba2610a04317bda192df2e53b71c"  # Replace with your Azure OpenAI API key 
openai.api_base = "https://genai-pssl-sweden.openai.azure.com/"  # Replace with your Azure OpenAI endpoint
openai.api_version = "2024-04-01-preview"
DEPLOYMENT_NAME = "pssl-gpt-4o"  # Replace with your deployed model name (e.g., "gpt-4o")

# === Sentence Embedding Model ===
embedder = SentenceTransformer('all-MiniLM-L6-v2')

# === Extract Q&A from Excel ===
def load_qa_from_excel(filepath):
    xls = pd.ExcelFile(filepath)
    chunks, meta_info = [], []
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        if 'Question' in df.columns and 'Answer' in df.columns:
            for i, row in df.iterrows():
                question = str(row.get("Question", "")).strip()
                answer = str(row.get("Answer", "")).strip()
                if question and answer:
                    chunk = f"Q: {question}\nA: {answer}"
                    chunks.append(chunk)
                    meta_info.append({'section': sheet, 'chunk_id': len(chunks)-1})
    return chunks, meta_info

# === Extract Text from SOW ===
def extract_text_from_sow(sow_filepath):
    doc = Document(sow_filepath)
    sow_text = [para.text for para in doc.paragraphs]
    return "\n".join(sow_text)

# === Build FAISS Vector Store ===
def build_faiss_index(chunks):
    embeddings = embedder.encode(chunks, convert_to_numpy=True)
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(embeddings)
    return index, embeddings

# === Extract Project Name ===
def extract_project_name(chunks):
    keywords = ["project name", "name of the project"]
    for chunk in chunks:
        lower_chunk = chunk.lower()
        if any(keyword in lower_chunk for keyword in keywords):
            for line in chunk.split("\n"):
                if line.lower().startswith("a:"):
                    name_candidate = line[2:].strip()
                    if len(name_candidate) <= 80:
                        return name_candidate
    return None

# === Add Cover Page ===
def add_cover_page(doc, project_name):
    doc.add_paragraph()
    title = doc.add_paragraph("BUSINESS REQUIREMENTS DOCUMENT")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].bold = True

    doc.add_paragraph()
    subtitle = doc.add_paragraph(project_name or "[Project Name Here]")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

# === Sections for Deep Attention ===
detailed_sections = {
    "Executive Summary",
    "Scope",
    "Functional Requirements",
    "Non-Functional Requirements",
    "KPIs",
    "Governance"
}

# === Prompt Template ===
def get_prompt_template(section, context):
    if section in {"Executive Summary", "Scope", "Objectives"}:
        return f"""
You are a top-tier Business Analyst and Prompt Engineer generating a **concise** version of the **{section}** section for a Business Requirements Document (BRD).

Instructions:
- Your audience is C-level executives and senior stakeholders.
- Deliver clear, crisp, high-impact content that fits on one short section — ideally **no more than 150 words**.
- Focus only on the most important facts or insights from the context.
- Avoid technical jargon, fluff, or repetition.
- Use short paragraphs or bullet points, but do not oversimplify.
- The language should be formal, strategic, and result-oriented.

Section: {section}
---------------------
Context:
{context}

Generate a **concise**, **executive-ready**, and **strategically framed** version of this section.
"""
    elif section in detailed_sections:
        return f"""
You are a senior business analyst generating the **{section}** section of a Business Requirements Document (BRD).

Instructions:
- This section is critical — provide detailed, structured, and exhaustive content using the context.
- Organize with logical subheadings.
- Break down the content into clear bullet points and numbered lists if applicable.
- Highlight any dependencies, metrics, or responsibilities explicitly.
- Avoid assumptions — strictly use the information provided in the context.

Section: {section}
---------------------
Context:
{context}

Generate a comprehensive, professional, and clearly structured version of this section.
"""
    else:
        return f"""
You are a senior business analyst generating the **{section}** section of a Business Requirements Document (BRD).

Instructions:
- Organize the content in a structured and logical format suitable for executives and stakeholders.
- Use clear subheadings where appropriate to break down the content.
- Use concise language and bullet points for readability.
- Do not include filler or generic text — base everything strictly on the context provided.
- The tone should be formal, informative, and boardroom-ready.

Section: {section}
---------------------
Context:
{context}

Please generate a neatly formatted and professional version of the above section.
"""

# === Generate with Azure OpenAI ===
def generate_section(prompt):
    response = openai.ChatCompletion.create(
        deployment_id=DEPLOYMENT_NAME,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a world-class Business Analyst and Prompt Engineer tasked with creating a highly professional, polished, "
                    "and boardroom-ready Business Requirements Document (BRD). Every section must be:\n"
                    "- Clear, concise, and formal\n"
                    "- Structured for C-level stakeholders\n"
                    "- Aligned with BRD best practices\n"
                    "- Grounded only in the provided context (no assumptions)\n"
                    "- Well-organized using headings, bullet points, and proper segmentation"
                )
            },
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=1024
    )
    return response['choices'][0]['message']['content'].strip()

# === Generate BRD Word File ===
def create_brd_doc(sections_content, project_name=None, output_path="Generated_BRD.docx"):
    doc = Document()
    add_cover_page(doc, project_name)
    doc.add_page_break()
    for section, content in sections_content.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    doc.save(output_path)
    return output_path

# === FastAPI App ===
app = FastAPI()

@app.post("/generate-brd/")
async def generate_brd(
    excel_file: UploadFile = File(...),
    sow_file: UploadFile = File(...)
):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_excel:
        shutil.copyfileobj(excel_file.file, temp_excel)
        excel_path = temp_excel.name

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_sow:
        shutil.copyfileobj(sow_file.file, temp_sow)
        sow_path = temp_sow.name

    chunks, metadata = load_qa_from_excel(excel_path)
    index, embeddings = build_faiss_index(chunks)
    sow_content = extract_text_from_sow(sow_path)
    project_name = extract_project_name(chunks)

    brd_sections = [
        "Executive Summary",
        "Scope",
        "Objectives",
        "Stakeholder Analysis",
        "Functional Requirements",
        "Non-Functional Requirements",
        "KPIs",
        "Governance",
        "Timeline and Milestones",
        "Risks and Mitigation"
    ]

    brd_result = {}
    for section in brd_sections:
        query_vec = embedder.encode([section])
        top_k = 15 if section in detailed_sections else 8
        _, indices = index.search(np.array(query_vec), k=top_k)
        retrieved = "\n\n".join([chunks[i] for i in indices[0]]) + "\n\n" + sow_content
        prompt = get_prompt_template(section, retrieved)
        response = generate_section(prompt)
        brd_result[section] = response

    output_path = os.path.join(tempfile.gettempdir(), "Generated_BRD.docx")
    create_brd_doc(brd_result, project_name=project_name, output_path=output_path)

    return FileResponse(output_path, filename="Generated_BRD.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
